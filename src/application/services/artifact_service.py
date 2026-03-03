import os
from typing import Dict, Any, List
import logging

logger = logging.getLogger(__name__)

class ArtifactService:
    def __init__(self, output_dir: str):
        self.output_dir = output_dir
        os.makedirs(self.output_dir, exist_ok=True)
        
    def generate_docx_report(self, filename: str, title: str, content: Dict[str, Any]) -> str:
        """
        Generate a DOCX report from structured content.
        content schema:
        {
            "summary": str,
            "evidence": List[Dict] (source, text),
            "compliance": str,
            "thinking": str
        }
        """
        try:
            from docx import Document
            from docx.shared import Pt, Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            logger.error("python-docx not installed. Cannot generate DOCX artifact.")
            return ""

        doc = Document()
        
        # Title
        head = doc.add_heading(title, 0)
        head.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Executive Summary
        doc.add_heading('Executive Summary', level=1)
        doc.add_paragraph(content.get('summary', 'No summary provided.'))
        
        # Compliance Status
        if content.get('compliance'):
            doc.add_heading('Compliance Status', level=1)
            doc.add_paragraph(content['compliance'])
            
        # Evidence Table
        if content.get('evidence'):
            doc.add_heading('Evidence & References', level=1)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Source Document'
            hdr_cells[1].text = 'Excerpt / Fact'
            
            for item in content['evidence']:
                row_cells = table.add_row().cells
                row_cells[0].text = item.get('source', 'Unknown')
                row_cells[1].text = item.get('text', '')

        # Thinking Process (optional / appendix)
        if content.get('thinking'):
            doc.add_page_break()
            doc.add_heading('Appendix A: Reasoning Process', level=1)
            doc.add_paragraph(content['thinking'])

        # Save
        full_path = os.path.join(self.output_dir, filename)
        if not full_path.lower().endswith('.docx'):
            full_path += '.docx'
            
        try:
            doc.save(full_path)
            logger.info(f"Artifact saved to {full_path}")
            return full_path
        except Exception as e:
            logger.error(f"Failed to save DOCX: {e}")
            return ""

    def generate_pdf_report(self, filename: str, title: str, content: Dict[str, Any]) -> str:
        """Generate a clean engineering PDF report."""
        try:
            from fpdf import FPDF
        except ImportError:
            logger.error("fpdf2 not installed.")
            return ""

        pdf = FPDF()
        pdf.add_page()
        
        # AECO Header
        pdf.set_fill_color(30, 30, 30)
        pdf.rect(0, 0, 210, 25, 'F')
        pdf.set_text_color(255, 255, 255)
        pdf.set_font("Arial", "B", 20)
        pdf.cell(0, 15, "SERAPEUM AECO INTELLIGENCE", ln=True, align="C")
        pdf.set_text_color(0, 0, 0)
        pdf.set_font("Arial", "", 8)
        pdf.cell(0, 5, "Generated Professional Engineering Deliverable", ln=True, align="C")
        pdf.ln(10)

        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, title, ln=True, align="L")
        pdf.ln(5)

        pdf.set_font("Arial", "B", 12)
        pdf.cell(0, 10, "Executive Summary", ln=True)
        pdf.set_font("Arial", "", 10)
        pdf.multi_cell(0, 5, content.get("summary", "No summary provided."))
        pdf.ln(5)

        if content.get("evidence"):
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 10, "Evidence & References", ln=True)
            pdf.set_font("Arial", "", 9)
            for item in content["evidence"]:
                pdf.multi_cell(0, 5, f"[{item.get('source')}]: {item.get('text')}")
                pdf.ln(2)

        full_path = os.path.join(self.output_dir, filename)
        if not full_path.lower().endswith(".pdf"): full_path += ".pdf"
        
        try:
            pdf.output(full_path)
            return full_path
        except Exception as e:
            logger.error(f"Failed to save PDF: {e}")
            return ""

    def generate_excel_boq(self, filename: str, data: List[Dict[str, Any]]) -> str:
        """Export material quantities/BOQ to Excel."""
        try:
            import pandas as pd
        except ImportError:
            logger.error("pandas not installed.")
            return ""

        full_path = os.path.join(self.output_dir, filename)
        if not full_path.lower().endswith(".xlsx"): full_path += ".xlsx"

        try:
            df = pd.DataFrame(data)
            df.to_excel(full_path, index=False)
            return full_path
        except Exception as e:
            logger.error(f"Failed to save Excel: {e}")
            return ""
