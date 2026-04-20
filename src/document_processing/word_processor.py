# -*- coding: utf-8 -*-
from __future__ import annotations

import os
import logging
from typing import Any, Dict, List

from .processor_utils import stable_doc_id
from src.domain.intelligence.text_chunker import TextChunker

logger = logging.getLogger(__name__)

MAX_DOC_CHARS = 500_000


class WordProcessor:
    """DOCX ingestion via python-docx (DOC is treated as unsupported text-only)."""

    def process(
        self,
        abs_path: str,
        rel_path: str,
        export_root: str,
        *,
        doc_id_override: str | None = None,
    ) -> Dict[str, Any]:
        doc_id = doc_id_override or stable_doc_id(abs_path, prefix="docx")
        ext = os.path.splitext(abs_path)[1].lower()

        text_parts: List[str] = []

        if ext == ".docx":
            try:
                from docx import Document  # python-docx

                d = Document(abs_path)

                # paragraphs
                for p in d.paragraphs:
                    t = (p.text or "").strip()
                    if t:
                        text_parts.append(t)

                # tables
                for table in d.tables:
                    text_parts.append("[Table]")
                    for row in table.rows:
                        cells = [(c.text or "").strip().replace('\n', ' ') for c in row.cells]
                        line = " | ".join(cells)
                        if line.strip('| '):
                            text_parts.append(line)
                            
                # inline images
                for rel in d.part.rels.values():
                    if "image" in rel.target_ref:
                        try:
                            image_part = rel.target_part
                            img_path = os.path.join(export_root, f"word_img_{rel.rId}.png")
                            with open(img_path, "wb") as f:
                                f.write(image_part.blob)
                                
                            from src.vision.ocr_backends import TesseractBackend
                            ocr_text = TesseractBackend().recognize(img_path, lang_hint="eng")
                            if ocr_text:
                                text_parts.append(f"[Image OCR] {ocr_text.strip()}")
                        except Exception as e:
                            logger.warning(f"Failed to extract word image: {e}")

            except Exception as e:
                logger.exception("   [WordProcessor] Failed to parse DOCX: %s", rel_path)
                text_parts.append(f"[docx] failed to parse: {e}")
        elif ext == ".doc":
            # .doc requires external conversion; keep a minimal placeholder
            text_parts.append("[doc] legacy .doc detected; conversion not implemented in this build.")
        else:
            text_parts.append(f"[word] unsupported extension: {ext}")

        # Deduplicate sequential identical lines + truncate protection
        final_parts: List[str] = []
        last_line: str | None = None
        current_len = 0

        for part in text_parts:
            s_part = part.strip()
            if not s_part:
                continue
            if s_part == last_line:
                continue

            final_parts.append(s_part)
            current_len += len(s_part)
            last_line = s_part

        text = "\n".join(final_parts).strip()

        # Blocks
        blocks = []
        if text:
            # Support either TextChunker.chunk_text(...) or TextChunker().chunk_text(...)
            try:
                blocks = TextChunker.chunk_text(text, source_type="docx")  # type: ignore[attr-defined]
            except Exception:
                try:
                    blocks = TextChunker().chunk_text(text, source_type="docx")  # type: ignore[call-arg]
                except Exception:
                    logger.exception("   [WordProcessor] TextChunker failed for %s", rel_path)
                    blocks = []

        return {
            "doc_id": doc_id,
            "text": text,
            "blocks": blocks,
            "pages": [
                {
                    "page_index": 0,
                    "py_text": text,
                    "text_hint": (text[:200] if text else rel_path),
                    "has_raster": 0,
                    "has_vector": 0,
                    "quality": "queued",
                }
            ],
            "meta": {"source": "word-processor", "rel_path": rel_path},
        }
